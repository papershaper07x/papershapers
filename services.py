# services.py
import os
import json
import logging
import asyncio
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
import concurrent.futures
from typing import Any, Dict, List

import numpy as np
import pandas as pd
from filelock import FileLock
from fastapi import BackgroundTasks, HTTPException, UploadFile, status
import fitz
# Third-party AI clients
import google.generativeai as genai

# --- New modular imports ---
import config
import models
import utils
import io
import tempfile
import shutil
import zipfile
from PIL import Image
import pandas as pd
from docx import Document


# Import specific functions from the original codebase structure
# These would ideally be in their own modules too, but for this refactoring, we keep them here.

from llm_researcher.search import tavily_search
from llm_researcher.utils import scrape_webpage
from llm_researcher.prompts import (
    generate_report_prompt,
    auto_agent_instructions,
    generate_search_queries_prompt,
)

from logger import log  # <--- ADD THIS IMPORT

from pydantic import ValidationError # <--- Make sure this is imported at the top

# -------- Module-level State --------
LOG = logging.getLogger("uvicorn.error")
_executor: ThreadPoolExecutor = None
df_content = pd.DataFrame()
df_prompt = pd.DataFrame()

# For LLM Researcher
llm_chat_model = None
embedding_client = None

background_task_status: Dict[str, Any] = {}
# -------- Custom Exceptions --------
class SchemaNotFoundError(Exception):
    """Custom exception for when a schema is not found."""

    pass


class GenerationError(Exception):
    """Custom exception for failures during the generation process."""

    def __init__(self, message, raw_output=None):
        super().__init__(message)
        self.raw_output = raw_output


# -------- Initialization and State Management --------
def set_executor(executor: ThreadPoolExecutor):
    """Receives the executor instance from main.py during startup."""
    global _executor
    _executor = executor



def get_data_status():
    """Returns the row counts of the loaded dataframes for the health check."""
    return len(df_content), len(df_prompt)


# -------- Core Service for /generate_full --------

# In services.py
# In services.py
from pydantic import ValidationError # Make sure this is imported
from copy import deepcopy # <--- ADD THIS IMPORT at the top of the file


# -------- Core Services for /process and /generate (Legacy) --------


async def handle_process_request(input_data: models.InputData):
    """Service logic for the /process endpoint."""
    if not input_data.is_logedIn:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED, detail="User not authenticated"
        )

    if input_data.answer:
        if not input_data.question_paper:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Missing 'question_paper' when answer is True",
            )
        response_text = await _get_answer(input_data)
    else:
        response_text = await _get_response(input_data)

    input_data.hit_count += 1
    _log_request(input_data, response_text)
    return response_text, input_data.hit_count


async def handle_legacy_generate(input_data: models.InputData):
    """Service logic for the legacy /generate endpoint."""
    if not input_data.is_logedIn:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED, detail="User not authenticated"
        )

    response_text = await _get_response(input_data)
    input_data.hit_count += 1
    _log_request(input_data, response_text)
    return response_text, input_data.hit_count


async def _get_response(data: models.InputData) -> str:
    """Internal function to generate a response based on CSV data."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(_executor, _get_response_sync, data)


def _get_response_sync(data: models.InputData) -> str:
    prompt_df_f = df_prompt[
        (df_prompt["Board"] == data.Board)
        & (df_prompt["Class"] == data.Class)
        & (df_prompt["Prompt_Type"] == data.Prompt_Type)
    ]
    content_df_f = df_content[
        (df_content["Board"] == data.Board)
        & (df_content["Class"] == data.Class)
        & (df_content["Subject"] == data.Subject)
        & (df_content["Chapter"] == data.Chapter)
    ]

    if content_df_f.empty or prompt_df_f.empty:
        raise HTTPException(
            status_code=404, detail="No matching records found for response generation"
        )

    prompt = prompt_df_f["Prompt_Data"].values[0]
    document_content = content_df_f["File_Data"].tolist()[0]

    msg = prompt.format(
        **{"Document_content": document_content, "Mock Paper": document_content}
    )
    return _generate_llm_response_text("models/gemini-2.5-flash-lite", msg)


async def _get_answer(data: models.InputData) -> str:
    """Internal function to generate an answer for a question paper."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(_executor, _get_answer_sync, data)


def _get_answer_sync(data: models.InputData) -> str:
    content_df_f = df_content[
        (df_content["Board"] == data.Board)
        & (df_content["Class"] == data.Class)
        & (df_content["Subject"] == data.Subject)
        & (df_content["Chapter"] == data.Chapter)
    ]
    if content_df_f.empty:
        raise HTTPException(
            status_code=404,
            detail="No matching content records found for answer generation",
        )

    document_content = content_df_f["File_Data"].tolist()[0]
    answer_paper_prompt = f"""You are a Expert teacher... Guideline for Question paper:- {data.question_paper}\n\nTextbook content:-{document_content}"""
    return _generate_llm_response_text(
        "models/gemini-2.5-flash-lite", answer_paper_prompt
    )


def _generate_llm_response_text(model: str, msg: str) -> str:
    """Helper to call the Gemini API and return text."""
    if not config.GOOGLE_API_KEY:
        raise RuntimeError("Google API key is not configured.")
    model_obj = genai.GenerativeModel(model)
    resp = model_obj.generate_content(msg)
    return resp.text


def _log_request(input_data: models.InputData, response_text: str):
    """Logs request and response details to a file with a lock."""
    log_file = Path(config.REQUEST_LOG_PATH)
    lock_file = str(log_file) + ".lock"
    log_file.parent.mkdir(parents=True, exist_ok=True)
    try:
        with FileLock(lock_file):
            with open(log_file, "a") as f:
                log_entry = input_data.dict()
                log_entry["response"] = response_text
                f.write(json.dumps(log_entry) + "\n")
    except Exception as e:
        LOG.error(f"Failed to write request log: {e}")


# -------- Core Service for /research --------


async def handle_research_request(query: str) -> str:
    """Service logic for the /research endpoint."""
    loop = asyncio.get_event_loop()
    # The entire research_conduct function is blocking, so run it in the executor
    report = await loop.run_in_executor(_executor, _research_conduct, query)
    return report


def _research_conduct(query: str) -> str:
    """The full, synchronous research pipeline. (Logic moved from main.py)"""
    LOG.info(f"Starting research for query: '{query}'")
    role_prompt = _choose_agent(query)
    search_results = tavily_search(query)
    sub_queries = _generate_sub_queries(query, search_results)

    if not isinstance(sub_queries, list):
        sub_queries = [sub_queries]
    sub_queries.append(query)

    subquery_url_map = _gather_urls_for_subqueries(sub_queries, config.NO_OF_SOURCEURLS)
    results = _process_subqueries_parallel(subquery_url_map)
    content_with_prompt = generate_report_prompt(query, results, "str")
    report = _llm_generate_report(role_prompt, content_with_prompt)
    LOG.info(f"Research finished for query: '{query}'")
    return report


# --- LLM Researcher Helper Functions (Internal to this service) ---


def _choose_agent(query: str) -> str:
    # ... (logic from choose_agent in main.py)
    messages = [
        {"role": "system", "content": auto_agent_instructions()},
        {"role": "user", "content": f"task: {query}"},
    ]
    combined = " ".join(f"{msg['role']}: {msg['content']}" for msg in messages)
    response = llm_chat_model.send_message(combined)
    return response.text


def _generate_sub_queries(query: str, context: Any) -> list:
    # ... (logic from generate_sub_queries in main.py)
    prompt = generate_search_queries_prompt(
        query, max_iterations=config.NO_OF_SUBQUERIES, context=context
    )
    messages = [
        {"role": "system", "content": prompt},
        {"role": "user", "content": f"task: {query}"},
    ]
    combined = " ".join(f"{msg['role']}: {msg['content']}" for msg in messages)
    response = llm_chat_model.send_message(combined)
    # The original code didn't parse this, assuming it's a list-like string. A more robust implementation would parse it.
    return response.text.split("\n") if response.text else []


# in services.py
def _gather_urls_for_subqueries(
    sub_queries: list, num_urls: int
) -> Dict[str, List[str]]:
    # ...
    with concurrent.futures.ThreadPoolExecutor() as tpe:
        results = list(
            tpe.map(lambda sq: tavily_search(sq, max_results=num_urls), sub_queries)
        )

    # This new logic checks that 'res' is a list and that each item in it is a
    # dictionary containing the 'url' key before trying to access it.
    url_map = {}
    for sq, res in zip(sub_queries, results):
        if isinstance(res, list):
            # Safely extract URLs from valid dictionary objects
            valid_urls = [
                item["url"] for item in res if isinstance(item, dict) and "url" in item
            ]
            if valid_urls:
                url_map[sq] = valid_urls
    return url_map


def _process_subqueries_parallel(subquery_url_map: Dict[str, List[str]]) -> str:
    # ... (logic from process_subqueries_parallel in main.py)
    tasks = [
        (url, subquery) for subquery, urls in subquery_url_map.items() for url in urls
    ]
    results = []
    with concurrent.futures.ThreadPoolExecutor() as tpe:
        future_to_task = {
            tpe.submit(_process_url, url, subquery): (url, subquery)
            for url, subquery in tasks
        }
        for future in concurrent.futures.as_completed(future_to_task):
            try:
                res = future.result()
                if res:
                    results.append(res)
            except Exception as e:
                LOG.error(f"Error processing URL {future_to_task[future][0]}: {e}")
    return "".join(results)


def _process_url(url: str, subquery: str) -> str:
    # ... (logic from process_url in main.py)
    try:
        content, _, _ = scrape_webpage(url)  # Assuming scrape_webpage is robust
        if content:
            compressed = _contextual_compression(content, subquery)
            return f"SOURCE: {url},\nRelevant Chunks: {compressed}\n\n"
    except Exception as e:
        LOG.warning(f"Could not process URL {url}: {e}")
    return ""


def _contextual_compression(content: str, query: str, k: int = 10) -> str:
    # ... (logic from ContextualCompression in main.py, now using the initialized client)
    content_chunks = utils.generate_tokens(content)
    if not content_chunks:
        return ""

    query_result = embedding_client.embed_content(
        model="models/text-embedding-004", content=[query]
    )
    query_emb = query_result["embedding"][0]

    chunk_result = embedding_client.embed_content(
        model="models/text-embedding-004", content=content_chunks
    )
    chunk_embeddings = [emb for emb in chunk_result["embedding"]]

    similarities = [
        utils.cosine_similarity(query_emb, chunk_emb) for chunk_emb in chunk_embeddings
    ]
    top_indices = np.argsort(similarities)[-k:][::-1]
    return "\n\n".join([content_chunks[i] for i in top_indices])


def _llm_generate_report(agent_role_prompt: str, content_with_prompt: str) -> str:
    # ... (logic from llm_generate_report in main.py)
    messages = [
        {"role": "system", "content": agent_role_prompt},
        {"role": "user", "content": content_with_prompt},
    ]
    combined = " ".join(f"{msg['role']}: {msg['content']}" for msg in messages)
    response = llm_chat_model.send_message(combined)
    return response.text






# =============================================================================
# CORE SERVICES FOR DOCUMENT PROCESSING (/upload-and-process)
# =============================================================================

def _process_with_gemini(content: Any, task_type: str) -> str:
    """Internal function to process content with Google Gemini."""
    try:
        model = genai.GenerativeModel(config.DOCUMENT_PROCESSING_MODEL_NAME)

        if task_type == "summarize_text":
            prompt = f"Please provide a concise summary of the following document:\n\n{content}"
            response = model.generate_content(prompt)
            return response.text
        elif task_type in ["summarize_images", "alt_text"]:
            prompt_parts = []
            if task_type == "summarize_images":
                prompt_parts.append("Provide a detailed summary of the document shown in the following page images.")
            else:
                prompt_parts.append("Provide a concise alt text and a brief one-sentence summary for the following image.")
            
            if isinstance(content, list):
                prompt_parts.extend(content)
            else:
                prompt_parts.append(content)
            
            response = model.generate_content(prompt_parts)
            return response.text
        else:
            raise ValueError("Invalid task type specified.")
    except Exception as e:
        log.error(f"Error with Google Gemini API: {e}")
        raise Exception(f"Error with Google Gemini API: {e}")


def _process_single_document(file_content_or_path, filename: str) -> str:
    """Internal function: Processes one document and returns its summary.

    Accepts either bytes (file_content) or a path to the file on disk. If a
    string path is provided, it will prefer streaming from disk to avoid extra
    memory usage.
    """
    is_path = isinstance(file_content_or_path, str)
    file_extension = os.path.splitext(filename)[1].lower()

    if file_extension == ".pdf":
        # Use utils.parse_pdf which now supports file_path to avoid loading bytes
        if is_path:
            # quick page count check via fitz without loading full bytes
            pdf_doc = fitz.open(file_content_or_path)
            if len(pdf_doc) > 10:
                raise ValueError(f"PDF '{filename}' exceeds the 10-page limit.")
            parsed_content = utils.parse_pdf(file_path=file_content_or_path, filename=filename)
        else:
            pdf_doc = fitz.open(stream=file_content_or_path, filetype="pdf")
            if len(pdf_doc) > 10:
                raise ValueError(f"PDF '{filename}' exceeds the 10-page limit.")
            parsed_content = utils.parse_pdf(file_content=file_content_or_path, filename=filename)

        if "text" in parsed_content:
            return _process_with_gemini(parsed_content["text"], "summarize_text")
        elif "images" in parsed_content:
            return _process_with_gemini(parsed_content["images"], "summarize_images")

    elif file_extension in [".jpg", ".jpeg", ".png"]:
        if is_path:
            image = Image.open(file_content_or_path)
        else:
            image = Image.open(io.BytesIO(file_content_or_path))
        return _process_with_gemini(image, "alt_text")

    elif file_extension == ".docx":
        if is_path:
            doc = Document(file_content_or_path)
        else:
            doc = Document(io.BytesIO(file_content_or_path))
        text = "\n".join([para.text for para in doc.paragraphs])
        return _process_with_gemini(text, "summarize_text")

    elif file_extension in [".xls", ".xlsx"]:
        if is_path:
            df = pd.read_excel(file_content_or_path)
        else:
            df = pd.read_excel(io.BytesIO(file_content_or_path))
        return _process_with_gemini(df.to_string(), "summarize_text")

    else:
        return f"Unsupported file type: {filename}"


def _process_file_background(file_content: bytes, filename: str, task_id: str):
    """Deprecated: keep for backwards compatibility. Prefer file-path based worker.

    This function accepts bytes and delegates to the path-based worker by
    writing the bytes to a temporary file and invoking that worker. That keeps
    behavior unchanged but avoids retaining bytes in long-lived memory.
    """
    tmp = None
    try:
        suffix = os.path.splitext(filename)[1] if getattr(filename, "__str__", None) else None
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        tmp.write(file_content)
        tmp.flush()
        tmp.close()
        _process_file_background_from_path(tmp.name, filename, task_id)
    finally:
        try:
            if tmp and os.path.exists(tmp.name):
                os.unlink(tmp.name)
        except Exception:
            pass


def _process_file_background_from_path(file_path: str, filename: str, task_id: str):
    """Background worker entrypoint which reads from a file on disk instead of
    accepting bytes. This prevents keeping large uploads in memory.
    The function delegates to the existing processing functions and removes
    temporary files when finished.
    """
    try:
        file_extension = os.path.splitext(filename)[1].lower()

        if file_extension == ".zip":
            all_summaries = []
            with zipfile.ZipFile(file_path) as zf:
                valid_files = [f for f in zf.infolist() if not f.is_dir() and not f.filename.startswith('__MACOSX/')]

                if len(valid_files) > 5:
                    raise ValueError(f"ZIP archive exceeds the 5-file limit.")

                for file_info in valid_files:
                    member_tmp_name = None
                    try:
                        with zf.open(file_info) as member_f:
                            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file_info.filename)[1]) as member_tmp:
                                shutil.copyfileobj(member_f, member_tmp)
                                member_tmp_name = member_tmp.name

                        summary = _process_single_document(member_tmp_name, file_info.filename)
                        all_summaries.append({"filename": file_info.filename, "summary": summary})
                    except Exception as doc_error:
                        log.error(f"Error processing {file_info.filename} in zip: {doc_error}")
                        all_summaries.append({"filename": file_info.filename, "summary": f"Could not process. Error: {doc_error}"})
                    finally:
                        try:
                            if member_tmp_name and os.path.exists(member_tmp_name):
                                os.unlink(member_tmp_name)
                        except Exception:
                            pass

            background_task_status[task_id] = {"summaries": all_summaries}
        else:
            summary = _process_single_document(file_path, filename)
            background_task_status[task_id] = {"summaries": [{"filename": filename, "summary": summary}]}

    except ValueError as ve:
        if "exceeds" in str(ve) and "limit" in str(ve):
            background_task_status[task_id] = {"requires_payment": True, "reason": str(ve)}
        else:
            background_task_status[task_id] = {"error": str(ve)}
    except Exception as e:
        log.error(f"Error processing file in background for task {task_id}: {e}")
        background_task_status[task_id] = {"error": str(e)}
    finally:
        # Cleanup the uploaded file (if it's a temp file path)
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
        except Exception:
            pass


# --- Handler functions to be called by endpoints ---

async def handle_upload_and_process(background_tasks: BackgroundTasks, file: Any):
    """Service handler for creating and dispatching a document processing task."""
    # Stream uploaded file to a temporary file to avoid holding its bytes in memory.
    suffix = os.path.splitext(file.filename)[1] if getattr(file, "filename", None) else None
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    try:
        await file.seek(0)
        while True:
            chunk = await file.read(65536)
            if not chunk:
                break
            tmp.write(chunk)
        tmp.flush()
        tmp.close()

        task_id = f"task_{file.filename}_{os.urandom(4).hex()}"
        background_task_status[task_id] = "processing"

        # Schedule the path-based worker which will remove the temp file when done
        background_tasks.add_task(_process_file_background_from_path, tmp.name, file.filename, task_id)
        return task_id
    except Exception:
        try:
            tmp.close()
        except Exception:
            pass
        if os.path.exists(tmp.name):
            os.unlink(tmp.name)
        raise


def handle_get_task_status(task_id: str):
    """Service handler for retrieving the status of a task."""
    status = background_task_status.get(task_id)
    if not status:
        raise HTTPException(status_code=404, detail="Task not found")
    return status






# In services.py, add this new import at the top
import json

# ... (other imports)


def _get_resume_analysis_prompt(analysis_type: models.AnalysisType) -> str:
    """
    Selects the expert-level system prompt that instructs the LLM to return a structured JSON object
    that perfectly matches the frontend's data contract.
    """
    base_persona = (
        "You are an expert career coach and professional resume reviewer acting as a data extraction API. "
        "Your task is to analyze the provided resume text and return a single, valid JSON object. "
        "Do NOT add any introductory text, explanations, or Markdown formatting like ```json. "
        "Your entire output must be only the JSON object."
    )

    # This JSON structure now perfectly matches the frontend's TypeScript interfaces.
    json_structure = """
    {
      "score": {
        "overall": <integer, 0-100, a holistic score>,
        "skills": <integer, 0-100, score for skills presentation and relevance>,
        "experience": <integer, 0-100, score for impact and quality of experience section>,
        "education": <integer, 0-100, score for clarity and relevance of education>
      },
      "personalInfo": {
        "name": "<string, extracted name or null>",
        "email": "<string, extracted email or null>",
        "phone": "<string, extracted phone number or null>",
        "location": "<string, extracted city/state or null>"
      },
      "summary": "<string, a 2-4 sentence professional summary based on the resume>",
      "skills": ["<string, skill 1>", "<string, skill 2>", "..."],
      "experience": [
        {
          "position": "<string, job title>",
          "company": "<string, company name>",
          "duration": "<string, e.g., 'Jan 2022 - Present'>",
          "description": "<string, a 1-2 sentence summary of the role's key responsibilities and achievements>"
        }
      ],
      "education": [
        {
          "degree": "<string, e.g., 'Bachelor of Science in Computer Science'>",
          "institution": "<string, university name>",
          "year": "<string, e.g., '2018 - 2022'>"
        }
      ],
      "recommendations": {
        "strengths": ["<string, a key strength of the resume>"],
        "improvements": ["<string, a critical area for improvement>"],
        "suggestions": ["<string, an actionable suggestion>"]
      }
    }
    """
    
    focus_instruction = {
        models.AnalysisType.GENERAL: "Provide a balanced, general analysis focusing on overall presentation and impact.",
        models.AnalysisType.DETAILED: "Provide a detailed analysis, paying close attention to every section. Be critical in your scoring.",
        models.AnalysisType.SKILLS: "Focus your analysis heavily on the skills section. Score the 'skills' field highest. Recommendations should be skills-focused.",
        models.AnalysisType.EXPERIENCE: "Focus your analysis on the work experience section. Score the 'experience' field highest. Recommendations should be experience-focused."
    }

    task_prompt = focus_instruction.get(analysis_type)

    return (
        f"{base_persona}\n\n"
        f"Analysis Focus: {task_prompt}\n\n"
        f"Based on the resume text provided, populate the following JSON structure. Ensure all fields are filled accurately. "
        f"If a piece of information is not present, use null for optional fields and empty arrays [] for lists.\n\n"
        f"JSON Structure to populate:\n{json_structure}"
    )


async def handle_resume_analysis(file: UploadFile, analysis_type: models.AnalysisType):
    """
    Service handler for the resume analysis workflow, now returns a structured dictionary.
    """
    if file.size > config.MAX_RESUME_FILE_SIZE:
        raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="File size exceeds 10MB.")

    try:
        # Stream upload to temp file to avoid holding all bytes in memory
        suffix = os.path.splitext(file.filename)[1] if getattr(file, "filename", None) else None
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        try:
            await file.seek(0)
            while True:
                chunk = await file.read(65536)
                if not chunk:
                    break
                tmp.write(chunk)
            tmp.flush()
            tmp.close()

            resume_text = utils.parse_document_to_text(file_path=tmp.name, filename=file.filename)
        finally:
            try:
                if tmp and os.path.exists(tmp.name):
                    os.unlink(tmp.name)
            except Exception:
                pass
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))

    prompt = _get_resume_analysis_prompt(analysis_type)
    full_prompt = f"{prompt}\n\n--- RESUME TEXT TO ANALYZE ---\n\n{resume_text}"

    try:
        log.info(f"Sending resume for '{analysis_type.value}' JSON analysis to Gemini.")
        model = genai.GenerativeModel(config.RESUME_ANALYSIS_MODEL_NAME)
        
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(
            _executor, 
            lambda: model.generate_content(full_prompt)
        )
        
        # --- CRITICAL CHANGE: Parse the LLM's text response as JSON ---
        try:
            # Clean the response text to remove potential markdown backticks
            cleaned_text = response.text.strip().replace("```json", "").replace("```", "")
            analysis_json = json.loads(cleaned_text)
            log.info("Successfully received and parsed JSON analysis from Gemini.")
            return analysis_json
        except json.JSONDecodeError:
            log.error(f"Failed to parse JSON from Gemini response. Raw response: {response.text}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="AI model returned a malformed response. Please try again."
            )

    except Exception as e:
        log.error(f"Error during Gemini API call for resume analysis: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="AI model failed to generate a response.")
    """
    Service handler for the entire resume analysis workflow.
    """
    # 1. Validate File Size
    if file.size > config.MAX_RESUME_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File size exceeds the limit of {config.MAX_RESUME_FILE_SIZE / (1024*1024)}MB."
        )

    # 2. Read and Parse File Content
    try:
        suffix = os.path.splitext(file.filename)[1] if getattr(file, "filename", None) else None
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        try:
            await file.seek(0)
            while True:
                chunk = await file.read(65536)
                if not chunk:
                    break
                tmp.write(chunk)
            tmp.flush()
            tmp.close()

            resume_text = utils.parse_document_to_text(file_path=tmp.name, filename=file.filename)
        finally:
            try:
                if tmp and os.path.exists(tmp.name):
                    os.unlink(tmp.name)
            except Exception:
                pass
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))

    # 3. Get the Expert Prompt
    prompt = _get_resume_analysis_prompt(analysis_type)
    full_prompt = f"{prompt}\n\n--- RESUME TEXT TO ANALYZE ---\n\n{resume_text}"

    # 4. Call the LLM for Analysis
    try:
        log.info(f"Sending resume for '{analysis_type.value}' analysis to Gemini.")
        model = genai.GenerativeModel(config.RESUME_ANALYSIS_MODEL_NAME)
        
        # Run the blocking network call in the thread pool executor
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(
            _executor, 
            lambda: model.generate_content(full_prompt)
        )
        
        log.info("Successfully received analysis from Gemini.")
        return response.text
    except Exception as e:
        log.error(f"Error during Gemini API call for resume analysis: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An error occurred while analyzing the resume with the AI model."
        )